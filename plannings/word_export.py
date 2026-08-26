import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE

MOIS_FR = ["", "Janvier", "Février", "Mars", "Avril", "Mai", "Juin",
           "Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"]

TITRES_PAR_TYPE = {
    "garde": "Tableau de garde du personnel",
    "permanence": "Tableau de permanence",
    "astreinte": "Tableau d'astreinte des médecins",
}


def exporter_planning_word(planning, jours, postes, gardes_par_cle, libelle_type,
                           service_nom, est_permanence=False, type_activite="garde"):
    """Planning A4 portrait sur une seule page.
       Haut et bas laissés entièrement vides pour logos et signatures."""
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    # Marges : espace vide en haut pour logos, en bas pour signatures
    section.top_margin = Cm(3.8)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # ── TITRES ──
    p_service = doc.add_paragraph()
    r_service = p_service.add_run(f"Service/unité : {service_nom}")
    r_service.bold = True
    r_service.font.size = Pt(11)
    p_service.paragraph_format.space_after = Pt(4)

    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titre = p_titre.add_run(TITRES_PAR_TYPE.get(type_activite, f"Tableau de {libelle_type.lower()}"))
    r_titre.bold = True
    r_titre.underline = True
    r_titre.font.size = Pt(15)
    p_titre.paragraph_format.space_after = Pt(4)

    p_mois = doc.add_paragraph()
    p_mois.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_mois = p_mois.add_run(f"Mois : {MOIS_FR[planning.mois]} {planning.annee}")
    r_mois.bold = True
    r_mois.font.size = Pt(11)
    p_mois.paragraph_format.space_after = Pt(8)

    # ── TABLEAU PRINCIPAL ──
    nb_cols = 1 + len(postes)
    table = doc.add_table(rows=1, cols=nb_cols)
    table.style = 'Table Grid'
    table.autofit = False

    entete = table.rows[0].cells
    entete[0].text = "Jour"
    for i, poste in enumerate(postes, start=1):
        if est_permanence:
            libelle_poste = f"{libelle_type} ({poste['heure_debut'][:5]}-{poste['heure_fin'][:5]})"
        else:
            libelle_poste = f"{libelle_type} {'de jour' if poste['type_vacation'] == 'jour' else 'de nuit'} ({poste['heure_debut'][:5]}-{poste['heure_fin'][:5]})"
        entete[i].text = libelle_poste

    for jour in jours:
        row = table.add_row().cells
        row[0].text = f"{jour['numero']} {jour['nom_jour']}"
        for i, poste in enumerate(postes, start=1):
            cle = f"{jour['numero']}-{poste['id']}"
            garde = gardes_par_cle.get(cle)
            if not garde:
                row[i].text = ""
            elif est_permanence:
                row[i].text = ", ".join(f"{a['agent_nom']} ({a['heures'] or 0:g}h)" for a in garde['affectations'])
            else:
                row[i].text = ", ".join(a['agent_nom'] for a in garde['affectations'])

    # Format agrandi pour remplir la page sans déborder
    for row in table.rows:
        row.height = Cm(0.58)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if not p.runs:
                    p.add_run("")
                for run in p.runs:
                    run.font.size = Pt(10)

    # Largeur des colonnes : utiliser toute la largeur utile (19 cm)
    usable_width = Cm(19.0)
    col_width = usable_width / nb_cols
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
